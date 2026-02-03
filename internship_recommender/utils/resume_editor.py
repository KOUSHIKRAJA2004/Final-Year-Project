"""
AI-Powered Resume Editor
Intelligently modifies resumes to match job descriptions while preserving layout
Uses Ollama for LLM-based text generation and python-docx for DOCX manipulation
"""
import re
import os
from typing import Dict, List, Tuple, Optional
from dataclasses import dataclass
import json

try:
    from docx import Document
    from docx.shared import RGBColor, Pt
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    Document = None

try:
    import requests
    REQUESTS_AVAILABLE = True
except ImportError:
    REQUESTS_AVAILABLE = False


@dataclass
class ResumeSection:
    """Represents a section of the resume"""
    name: str
    content: str
    start_index: int
    end_index: int


@dataclass
class ModificationSuggestion:
    """Represents a suggested modification"""
    section: str
    original_text: str
    suggested_text: str
    reason: str
    skill_added: Optional[str] = None
    impact_score: float = 0.0  # Expected ATS score increase


class ResumeEditor:
    """
    AI-powered resume editor using Ollama
    """
    
    def __init__(self, ollama_url: str = "http://localhost:11434", model: str = "tinyllama"):
        """
        Initialize the resume editor
        
        Args:
            ollama_url: URL of Ollama API
            model: Model to use (llama2, mistral, etc.)
        """
        self.ollama_url = ollama_url
        self.model = model
        self.ollama_available = self._check_ollama()
    
    def _check_ollama(self) -> bool:
        """Check if Ollama is running"""
        if not REQUESTS_AVAILABLE:
            print("Warning: requests library not available")
            return False
        
        try:
            response = requests.get(f"{self.ollama_url}/api/tags", timeout=2)
            return response.status_code == 200
        except Exception as e:
            print(f"Warning: Ollama not available at {self.ollama_url}: {e}")
            return False
    
    def _call_ollama(self, prompt: str, system_prompt: str = None) -> str:
        """
        Call Ollama API to generate text
        
        Args:
            prompt: User prompt
            system_prompt: System/instruction prompt
            
        Returns:
            Generated text
        """
        if not self.ollama_available:
            return "[Ollama not available - please start Ollama service]"
        
        try:
            payload = {
                "model": self.model,
                "prompt": prompt,
                "stream": False
            }
            
            if system_prompt:
                payload["system"] = system_prompt
            
            response = requests.post(
                f"{self.ollama_url}/api/generate",
                json=payload,
                timeout=30
            )
            
            if response.status_code == 200:
                result = response.json()
                return result.get("response", "")
            else:
                return f"[Ollama error: {response.status_code}]"
        
        except Exception as e:
            print(f"Error calling Ollama: {e}")
            return f"[Error: {str(e)}]"
    
    def extract_sections(self, resume_text: str) -> List[ResumeSection]:
        """
        Extract sections from resume text
        
        Common sections: Summary, Experience, Education, Skills, Projects, Certifications
        """
        sections = []
        
        # Common section headers (case-insensitive)
        section_patterns = [
            r'(SUMMARY|PROFILE|OBJECTIVE)',
            r'(EXPERIENCE|WORK HISTORY|EMPLOYMENT)',
            r'(EDUCATION|ACADEMIC)',
            r'(SKILLS|TECHNICAL SKILLS|CORE COMPETENCIES)',
            r'(PROJECTS|KEY PROJECTS)',
            r'(CERTIFICATIONS|CERTIFICATES)',
            r'(ACHIEVEMENTS|ACCOMPLISHMENTS)',
        ]
        
        lines = resume_text.split('\n')
        current_section = None
        current_content = []
        current_start = 0
        
        for i, line in enumerate(lines):
            line_upper = line.strip().upper()
            
            # Check if line is a section header
            is_header = False
            for pattern in section_patterns:
                if re.match(pattern, line_upper):
                    # Save previous section
                    if current_section:
                        sections.append(ResumeSection(
                            name=current_section,
                            content='\n'.join(current_content),
                            start_index=current_start,
                            end_index=i
                        ))
                    
                    current_section = line.strip()
                    current_content = []
                    current_start = i
                    is_header = True
                    break
            
            if not is_header and current_section:
                current_content.append(line)
        
        # Add last section
        if current_section:
            sections.append(ResumeSection(
                name=current_section,
                content='\n'.join(current_content),
                start_index=current_start,
                end_index=len(lines)
            ))
        
        return sections
    
    def suggest_resume_modifications(
        self,
        resume_text: str,
        jd_text: str,
        missing_skills: List[str],
        weak_skills: List[str],
        ats_score: float
    ) -> List[ModificationSuggestion]:
        """
        Generate intelligent modification suggestions using Ollama
        
        Args:
            resume_text: Original resume text
            jd_text: Job description text
            missing_skills: Skills missing from resume
            weak_skills: Skills that are weak matches
            ats_score: Current ATS score
            
        Returns:
            List of modification suggestions
        """
        suggestions = []
        
        # Extract sections
        sections = self.extract_sections(resume_text)
        
        # 1. Skills Section Enhancement
        skills_section = next((s for s in sections if 'SKILL' in s.name.upper()), None)
        if skills_section and missing_skills:
            suggestion = self._enhance_skills_section(
                skills_section.content,
                missing_skills,
                jd_text
            )
            if suggestion:
                suggestions.append(suggestion)
        
        # 2. Experience Section Enhancement
        exp_section = next((s for s in sections if 'EXPERIENCE' in s.name.upper()), None)
        if exp_section and (missing_skills or weak_skills):
            suggestion = self._enhance_experience_section(
                exp_section.content,
                missing_skills + weak_skills,
                jd_text
            )
            if suggestion:
                suggestions.append(suggestion)
        
        # 3. Summary/Objective Enhancement
        summary_section = next((s for s in sections if any(kw in s.name.upper() for kw in ['SUMMARY', 'PROFILE', 'OBJECTIVE'])), None)
        if summary_section:
            suggestion = self._enhance_summary_section(
                summary_section.content,
                jd_text,
                missing_skills
            )
            if suggestion:
                suggestions.append(suggestion)
        
        return suggestions
    
    def _enhance_skills_section(
        self,
        skills_content: str,
        missing_skills: List[str],
        jd_text: str
    ) -> Optional[ModificationSuggestion]:
        """Enhance skills section by adding missing skills"""
        
        if not missing_skills:
            return None
        
        # Generate prompt for Ollama
        system_prompt = """You are an expert professional resume writer. 
        Your task is to enhance the Skills section by naturally integrating missing skills. 
        Keep the format consistent with the original.
        CRITICAL: Your output MUST be in ENGLISH. Do not use any other language."""
        
        prompt = f"""Original Skills Section:
{skills_content}

Missing Skills to Add: {', '.join(missing_skills[:5])}

Job Description Context:
{jd_text[:500]}

Rewrite the Skills section to include the missing skills naturally. Maintain the original format and style.
Only output the enhanced Skills section, nothing else. Ensure all text is in English."""

        enhanced_text = self._call_ollama(prompt, system_prompt)
        
        if enhanced_text and not enhanced_text.startswith('['):
            return ModificationSuggestion(
                section="Skills",
                original_text=skills_content,
                suggested_text=enhanced_text.strip(),
                reason=f"Added {len(missing_skills[:5])} missing skills: {', '.join(missing_skills[:5])}",
                skill_added=', '.join(missing_skills[:5]),
                impact_score=len(missing_skills[:5]) * 3.0
            )
        
        return None
    
    def _enhance_experience_section(
        self,
        experience_content: str,
        target_skills: List[str],
        jd_text: str
    ) -> Optional[ModificationSuggestion]:
        """Enhance experience section to highlight target skills"""
        
        if not target_skills:
            return None
        
        system_prompt = """You are an expert professional resume writer. Enhance the Experience section by 
        rephrasing bullet points to highlight relevant skills from the job description. 
        Keep it truthful and maintain the original achievements.
        CRITICAL: Your output MUST be in ENGLISH. Do not use any other language."""
        
        prompt = f"""Original Experience Section:
{experience_content[:1000]}

Skills to Highlight: {', '.join(target_skills[:5])}

Job Description Keywords:
{jd_text[:500]}

Rewrite the experience bullet points to better highlight these skills. Keep the same roles and timeframes.
Only enhance the descriptions to better match the job requirements. Output only the enhanced section.
Ensure the output is entirely in ENGLISH."""

        enhanced_text = self._call_ollama(prompt, system_prompt)
        
        if enhanced_text and not enhanced_text.startswith('['):
            return ModificationSuggestion(
                section="Experience",
                original_text=experience_content[:500] + "...",
                suggested_text=enhanced_text.strip()[:500] + "...",
                reason=f"Rephrased to highlight: {', '.join(target_skills[:3])}",
                impact_score=5.0
            )
        
        return None
    
    def _enhance_summary_section(
        self,
        summary_content: str,
        jd_text: str,
        missing_skills: List[str]
    ) -> Optional[ModificationSuggestion]:
        """Enhance summary/objective to align with job description"""
        
        system_prompt = """You are an expert professional resume writer. Rewrite the professional summary to 
        align with the job description while staying truthful to the candidate's background.
        CRITICAL: Your output MUST be in ENGLISH. Do not use any other language."""
        
        prompt = f"""Original Summary:
{summary_content}

Job Description:
{jd_text[:400]}

Key Skills to Mention: {', '.join(missing_skills[:3])}

Rewrite the summary to better align with the job requirements. Output only the enhanced summary, 2-3 sentences max.
Ensure the output is entirely in ENGLISH."""

        enhanced_text = self._call_ollama(prompt, system_prompt)
        
        if enhanced_text and not enhanced_text.startswith('['):
            return ModificationSuggestion(
                section="Summary",
                original_text=summary_content,
                suggested_text=enhanced_text.strip(),
                reason="Aligned summary with job requirements",
                impact_score=4.0
            )
        
        return None

    def rewrite_text_segment(self, text: str, style: str = "professional") -> List[str]:
        """
        Rewrite a text segment in different styles or improved clarity.
        Returns a list of 3 variations.
        """
        system_prompt = f"""You are an expert professional resume writer. 
        Your task is to rewrite the provided text to be more {style}, impactful, and concise.
        Use strong action verbs.
        Output exactly 3 distinct variations labeled 1., 2., and 3.
        Do not include any conversational text or introductions. Just the numbered list.
        CRITICAL: Your output MUST be in ENGLISH. Do not use any other language."""
        
        prompt = f"Rewrite this resume content:\n'{text}'\n\nEnsure the output is 100% in English."
        
        response = self._call_ollama(prompt, system_prompt)
        
        # Parse response into list
        variations = []
        if response:
            # Simple parsing for numbered list
            lines = response.split('\n')
            for line in lines:
                clean_line = re.sub(r'^\d+[\.)]\s*', '', line).strip()
                if clean_line and len(clean_line) > 5:
                    variations.append(clean_line)
                    
        # Fallback if parsing fails or valid response is just text
        if not variations and response and not response.startswith('['):
            variations = [response.strip()]
            
        return variations[:3]

    def apply_modifications_to_docx(
        self,
        docx_path: str,
        suggestions: List[ModificationSuggestion],
        output_path: str
    ) -> bool:
        """
        Apply modifications to a DOCX file
        
        Args:
            docx_path: Path to original DOCX
            suggestions: List of approved modifications
            output_path: Path to save modified DOCX
            
        Returns:
            True if successful
        """
        if not DOCX_AVAILABLE:
            print("Error: python-docx not installed")
            return False
        
        try:
            # Load document
            doc = Document(docx_path)
            
            # Apply modifications section by section
            for suggestion in suggestions:
                self._apply_section_modification(doc, suggestion)
            
            # Save modified document
            doc.save(output_path)
            return True
        
        except Exception as e:
            print(f"Error applying modifications: {e}")
            return False
    
    def _apply_section_modification(self, doc: Document, suggestion: ModificationSuggestion):
        """Apply a single modification to the document"""
        
        # Find paragraphs containing the section name
        section_found = False
        for i, para in enumerate(doc.paragraphs):
            if suggestion.section.upper() in para.text.upper():
                section_found = True
                # Replace content in next paragraphs until next section
                j = i + 1
                content_replaced = False
                while j < len(doc.paragraphs):
                    # Check if we've hit another section header
                    if doc.paragraphs[j].text.strip().isupper() and len(doc.paragraphs[j].text.strip()) < 30:
                        break
                    
                    if not content_replaced:
                        # Replace first paragraph with new content
                        doc.paragraphs[j].text = suggestion.suggested_text
                        content_replaced = True
                    else:
                        # Clear subsequent paragraphs of this section
                        doc.paragraphs[j].text = ""
                    
                    j += 1
                break
        
        if not section_found:
            # If section not found, append at end
            doc.add_paragraph(f"\n{suggestion.section}", style='Heading 2')
            doc.add_paragraph(suggestion.suggested_text)


# Global instance
resume_editor = ResumeEditor()
